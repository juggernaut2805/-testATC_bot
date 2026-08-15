import os
import re
import random
import asyncio
import html

from aiohttp import web
from docx import Document

from aiogram import Bot, Dispatcher, F
from aiogram.filters import Command
from aiogram.types import CallbackQuery
from aiogram.utils.keyboard import InlineKeyboardBuilder
from aiogram.webhook.aiohttp_server import SimpleRequestHandler, setup_application


# ============================================================
# НАСТРОЙКИ
# ============================================================

TOKEN = os.getenv("BOT_TOKEN")
PORT = int(os.getenv("PORT", "10000"))
RENDER_URL = os.getenv("RENDER_EXTERNAL_URL")

FILE_NAME = "ТЕСТЫ.docx"

QUESTIONS_PER_TEST = 30
TEST_MINUTES = 30


if not TOKEN:
    raise RuntimeError("❌ Переменная BOT_TOKEN не найдена")

if not RENDER_URL:
    raise RuntimeError("❌ Переменная RENDER_EXTERNAL_URL не найдена")


# ============================================================
# BOT
# ============================================================

bot = Bot(TOKEN)
dp = Dispatcher()


# Все предметы
TESTS = {}

# Активные тесты пользователей
USERS = {}

# Таймеры
TIMERS = {}


# ============================================================
# ОЧИСТКА ТЕКСТА
# ============================================================

def clean(text):
    if not text:
        return ""

    text = text.replace("\xa0", " ")
    text = text.replace("\r", " ")
    text = text.replace("\n", " ")

    text = re.sub(r"\s+", " ", text)

    return text.strip()


# ============================================================
# ЭКРАНИРОВАНИЕ HTML
# ============================================================

def esc(text):
    return html.escape(str(text))


# ============================================================
# ЧТЕНИЕ DOCX
# ============================================================

def read_docx():

    if not os.path.exists(FILE_NAME):
        raise RuntimeError(
            f"❌ Файл {FILE_NAME} не найден"
        )

    doc = Document(FILE_NAME)

    lines = []

    # Обычные абзацы
    for paragraph in doc.paragraphs:

        text = paragraph.text.strip()

        if text:
            lines.append(text)

    # Если часть вопросов находится в таблицах
    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    text = paragraph.text.strip()

                    if text:
                        lines.append(text)

    return lines


# ============================================================
# ПРОВЕРКА: ЯВЛЯЕТСЯ ЛИ СТРОКА ВОПРОСОМ
# ============================================================

def is_question_start(line):

    return re.match(
        r"^Задание\s*#\s*\d+",
        line,
        re.IGNORECASE
    )


# ============================================================
# ПРОВЕРКА: ЯВЛЯЕТСЯ ЛИ СТРОКА ПРЕДМЕТОМ
# ============================================================

def get_subject(line):

    match = re.match(
        r"^\s*Тест\s*:\s*(.+?)\s*$",
        line,
        re.IGNORECASE
    )

    if match:
        return clean(match.group(1))

    return None


# ============================================================
# РАЗБОР ВАРИАНТА
#
# Поддерживает:
#
# 1) - текст
# 2) + текст
#
# И:
#
# 1) -
# текст
#
# 2) +
# текст
#
# И даже:
#
# 1)
# -
# текст
# ============================================================

def parse_option_line(line):

    # Вариант с номером и знаком
    match = re.match(
        r"^\s*(\d+)\)\s*([+-])\s*(.*)$",
        line
    )

    if match:

        return {
            "number": int(match.group(1)),
            "correct": match.group(2) == "+",
            "text": clean(match.group(3))
        }


    # Вариант только с номером
    match = re.match(
        r"^\s*(\d+)\)\s*$",
        line
    )

    if match:

        return {
            "number": int(match.group(1)),
            "correct": None,
            "text": ""
        }


    return None


# ============================================================
# ПАРСЕР
# ============================================================

def parse_tests():

    lines = read_docx()

    subjects = {}

    current_subject = None
    current_question = None
    current_option = None


    def save_option():

        nonlocal current_option
        nonlocal current_question

        if current_option is None:
            return

        if current_question is None:
            current_option = None
            return

        text = clean(
            current_option["text"]
        )

        if not text:
            current_option = None
            return

        option = {
            "number": current_option["number"],
            "text": text,
            "correct": bool(
                current_option["correct"]
            )
        }

        current_question["options"].append(
            option
        )

        if option["correct"]:
            current_question["correct"] = (
                option["number"]
            )

        current_option = None


    def save_question():

        nonlocal current_question

        save_option()

        if current_question is None:
            return

        if current_subject is None:
            current_question = None
            return

        question_text = clean(
            current_question["question"]
        )

        options = current_question["options"]

        correct = current_question["correct"]


        # Нам нужны нормальные вопросы
        if (
            question_text
            and len(options) >= 2
            and correct is not None
        ):

            subjects.setdefault(
                current_subject,
                []
            ).append({

                "number":
                    current_question["number"],

                "question":
                    question_text,

                "options":
                    options,

                "correct":
                    correct
            })


        current_question = None


    # ========================================================
    # ОСНОВНОЙ ЦИКЛ
    # ========================================================

    for raw_line in lines:

        line = clean(raw_line)

        if not line:
            continue


        # ----------------------------------------------------
        # НОВЫЙ ПРЕДМЕТ
        # ----------------------------------------------------

        subject = get_subject(line)

        if subject:

            save_question()

            current_subject = subject

            subjects.setdefault(
                current_subject,
                []
            )

            continue


        # ----------------------------------------------------
        # НОВЫЙ ВОПРОС
        # ----------------------------------------------------

        match = re.match(
            r"^Задание\s*#\s*(\d+)",
            line,
            re.IGNORECASE
        )

        if match:

            save_question()

            current_question = {

                "number":
                    int(match.group(1)),

                "question":
                    "",

                "options":
                    [],

                "correct":
                    None
            }

            current_option = None

            continue


        # ----------------------------------------------------
        # СЛУЖЕБНЫЙ ТЕКСТ
        # ----------------------------------------------------

        if line.lower().startswith(
            "выберите один"
        ):
            continue


        # ----------------------------------------------------
        # НОВЫЙ ВАРИАНТ
        # ----------------------------------------------------

        option = parse_option_line(line)

        if (
            option is not None
            and current_question is not None
        ):

            save_option()

            current_option = option

            continue


        # ----------------------------------------------------
        # Если вариант уже существует
        # ----------------------------------------------------

        if (
            current_option is not None
            and current_option["correct"] is None
            and not current_option["text"]
        ):

            # Например отдельной строкой стоит "+"
            if line == "+":
                current_option["correct"] = True
                continue

            # Или "-"
            if line == "-":
                current_option["correct"] = False
                continue


        # ----------------------------------------------------
        # Продолжение текста варианта
        # ----------------------------------------------------

        if current_option is not None:

            if current_option["text"]:

                current_option["text"] += (
                    " " + line
                )

            else:

                current_option["text"] = line

            continue


        # ----------------------------------------------------
        # Продолжение вопроса
        # ----------------------------------------------------

        if current_question is not None:

            if current_question["question"]:

                current_question["question"] += (
                    " " + line
                )

            else:

                current_question["question"] = line


    # Сохраняем последний вопрос
    save_question()


    return subjects


# ============================================================
# ЗАГРУЗКА ТЕСТОВ
# ============================================================

TESTS = parse_tests()


# ============================================================
# ЛОГИ RENDER
# ============================================================

print("")
print("============================================")
print("✈️ ATC TEST BOT")
print("============================================")

total_questions = 0

for subject, questions in TESTS.items():

    print(
        f"📚 {subject} -> "
        f"{len(questions)} вопросов"
    )

    total_questions += len(questions)

print("--------------------------------------------")

print(
    f"ВСЕГО ВОПРОСОВ: {total_questions}"
)

print("============================================")
print("")


# ============================================================
# КНОПКИ ПРЕДМЕТОВ
# ============================================================

def subjects_keyboard():

    keyboard = InlineKeyboardBuilder()

    for index, subject in enumerate(
        TESTS.keys()
    ):

        count = len(
            TESTS[subject]
        )

        keyboard.button(
            text=(
                f"📚 {subject} "
                f"({count})"
            ),
            callback_data=(
                f"subject:{index}"
            )
        )

    keyboard.adjust(1)

    return keyboard.as_markup()


# ============================================================
# /START
# ============================================================

@dp.message(Command("start"))
async def start(message):

    if not TESTS:

        await message.answer(
            "❌ Я не нашёл тесты в файле."
        )

        return


    text = (
        "✈️ <b>ATC TEST TRAINER</b>\n\n"
        "Выбери предмет:\n\n"
        "📝 30 вопросов\n"
        "⏱ 30 минут\n"
        "🔀 Вопросы перемешиваются\n"
        "🎲 Варианты перемешиваются\n"
        "📊 Результат и ответы — в конце\n\n"
        "Количество вопросов в каждом "
        "предмете указано в скобках."
    )


    await message.answer(
        text,
        parse_mode="HTML",
        reply_markup=subjects_keyboard()
    )


# ============================================================
# ВЫБОР ПРЕДМЕТА
# ============================================================

@dp.callback_query(
    F.data.startswith("subject:")
)
async def choose_subject(
    callback: CallbackQuery
):

    user_id = callback.from_user.id

    index = int(
        callback.data.split(":")[1]
    )

    subjects = list(
        TESTS.keys()
    )

    if index < 0 or index >= len(subjects):

        await callback.answer(
            "❌ Ошибка"
        )

        return


    subject = subjects[index]

    all_questions = TESTS[subject]

    count = len(all_questions)


    print(
        f"Пользователь {user_id}: "
        f"{subject} -> {count} вопросов"
    )


    # Проверяем минимум
    if count < QUESTIONS_PER_TEST:

        await callback.message.answer(
            "❌ В этом предмете найдено "
            f"<b>{count}</b> вопросов.\n\n"
            f"Для теста нужно минимум "
            f"<b>{QUESTIONS_PER_TEST}</b>.",
            parse_mode="HTML"
        )

        await callback.answer()

        return


    # Если старый тест есть
    if user_id in USERS:

        USERS.pop(
            user_id,
            None
        )


    # Останавливаем старый таймер
    old_timer = TIMERS.pop(
        user_id,
        None
    )

    if old_timer:

        old_timer.cancel()


    # ========================================================
    # БЕРЁМ 30 СЛУЧАЙНЫХ ВОПРОСОВ
    # ========================================================

    selected_questions = random.sample(
        all_questions,
        QUESTIONS_PER_TEST
    )


    test_questions = []


    for original in selected_questions:

        options = [
            {
                "number":
                    option["number"],

                "text":
                    option["text"],

                "correct":
                    option["correct"]
            }

            for option in original["options"]
        ]


        # Перемешиваем варианты
        random.shuffle(options)


        test_questions.append({

            "number":
                original["number"],

            "question":
                original["question"],

            "options":
                options,

            "correct":
                original["correct"]
        })


    USERS[user_id] = {

        "subject":
            subject,

        "questions":
            test_questions,

        "current":
            0,

        "answers":
            [],

        "finished":
            False
    }


    await callback.message.answer(
        "🚀 <b>ТЕСТ НАЧАЛСЯ!</b>\n\n"
        f"📚 {esc(subject)}\n"
        "📝 30 вопросов\n"
        "⏱ 30 минут\n\n"
        "Удачи! ✈️",
        parse_mode="HTML"
    )


    # Запускаем таймер
    TIMERS[user_id] = asyncio.create_task(
        test_timer(user_id)
    )


    await send_question(user_id)

    await callback.answer()


# ============================================================
# ТАЙМЕР
# ============================================================

async def test_timer(user_id):

    try:

        await asyncio.sleep(
            TEST_MINUTES * 60
        )

        if user_id in USERS:

            await finish_test(
                user_id,
                timeout=True
            )

    except asyncio.CancelledError:

        pass


# ============================================================
# ОТПРАВКА ВОПРОСА
# ============================================================

async def send_question(user_id):

    if user_id not in USERS:
        return


    user = USERS[user_id]


    if user["finished"]:
        return


    current = user["current"]

    questions = user["questions"]


    if current >= len(questions):

        await finish_test(
            user_id
        )

        return


    question = questions[current]


    text = (
        f"📝 <b>Вопрос {current + 1}/30</b>\n\n"
        f"{esc(question['question'])}"
    )


    keyboard = InlineKeyboardBuilder()


    for option in question["options"]:

        keyboard.button(
            text=(
                f"{option['number']}) "
                f"{option['text']}"
            ),
            callback_data=(
                f"answer:{option['number']}"
            )
        )


    keyboard.adjust(1)


    await bot.send_message(
        user_id,
        text,
        parse_mode="HTML",
        reply_markup=keyboard.as_markup()
    )


# ============================================================
# ОТВЕТ ПОЛЬЗОВАТЕЛЯ
# ============================================================

@dp.callback_query(
    F.data.startswith("answer:")
)
async def answer_question(
    callback: CallbackQuery
):

    user_id = callback.from_user.id


    if user_id not in USERS:

        await callback.answer(
            "Тест уже завершён."
        )

        return


    user = USERS[user_id]


    if user["finished"]:

        await callback.answer(
            "Тест уже завершён."
        )

        return


    selected_number = int(
        callback.data.split(":")[1]
    )


    current = user["current"]

    question = user["questions"][current]


    # Запоминаем ответ
    user["answers"].append({

        "question":
            question["question"],

        "selected":
            selected_number,

        "correct":
            question["correct"],

        "options":
            question["options"],

        "original_number":
            question["number"]
    })


    user["current"] += 1


    # Убираем старую кнопку
    try:

        await callback.message.delete()

    except Exception:

        pass


    # Последний вопрос
    if user["current"] >= 30:

        await finish_test(
            user_id
        )

    else:

        await send_question(
            user_id
        )


    await callback.answer()


# ============================================================
# ПОЛУЧИТЬ ТЕКСТ ВАРИАНТА
# ============================================================

def option_text(options, number):

    for option in options:

        if option["number"] == number:

            return option["text"]

    return "—"


# ============================================================
# РЕЗУЛЬТАТ
# ============================================================

async def finish_test(
    user_id,
    timeout=False
):

    if user_id not in USERS:
        return


    user = USERS[user_id]


    if user["finished"]:
        return


    user["finished"] = True


    # Останавливаем таймер
    timer = TIMERS.pop(
        user_id,
        None
    )

    if timer:

        timer.cancel()


    answers = user["answers"]


    correct_count = 0


    for answer in answers:

        if (
            answer["selected"]
            == answer["correct"]
        ):

            correct_count += 1


    answered = len(answers)

    wrong_count = (
        answered - correct_count
    )


    percent = 0

    if answered:

        percent = round(
            correct_count
            / answered
            * 100
        )


    if timeout:

        title = (
            "⏰ <b>ВРЕМЯ ВЫШЛО!</b>"
        )

    else:

        title = (
            "🏁 <b>ТЕСТ ЗАВЕРШЁН!</b>"
        )


    result = (
        f"{title}\n\n"
        f"📚 <b>{esc(user['subject'])}</b>\n\n"
        f"📝 Отвечено: "
        f"<b>{answered}/30</b>\n"
        f"✅ Правильно: "
        f"<b>{correct_count}</b>\n"
        f"❌ Ошибок: "
        f"<b>{wrong_count}</b>\n"
        f"📊 Результат: "
        f"<b>{percent}%</b>\n"
    )


    await bot.send_message(
        user_id,
        result,
        parse_mode="HTML"
    )


    # ========================================================
    # ОТВЕТЫ
    # ========================================================

    answer_text = (
        "📋 <b>ПРАВИЛЬНЫЕ ОТВЕТЫ</b>\n\n"
    )


    for i, answer in enumerate(
        answers,
        1
    ):

        correct_number = answer["correct"]

        correct_text = option_text(
            answer["options"],
            correct_number
        )


        selected_number = answer["selected"]

        selected_text = option_text(
            answer["options"],
            selected_number
        )


        if selected_number == correct_number:

            mark = "✅"

        else:

            mark = "❌"


        block = (
            f"<b>{i}. "
            f"Вопрос №{answer['original_number']}</b>\n"
            f"Правильный: "
            f"{correct_number}) "
            f"{esc(correct_text)}\n"
            f"Ваш ответ: "
            f"{selected_number}) "
            f"{esc(selected_text)} "
            f"{mark}\n\n"
        )


        # Telegram ограничивает длину сообщения
        if len(answer_text) + len(block) > 3800:

            await bot.send_message(
                user_id,
                answer_text,
                parse_mode="HTML"
            )

            answer_text = ""


        answer_text += block


    if answer_text:

        await bot.send_message(
            user_id,
            answer_text,
            parse_mode="HTML"
        )


    await bot.send_message(
        user_id,
        "🔄 Чтобы пройти новый тест, "
        "нажми /start"
    )


    USERS.pop(
        user_id,
        None
    )


# ============================================================
# WEBHOOK
# ============================================================

async def on_startup(app):

    webhook_url = (
        RENDER_URL.rstrip("/")
        + "/webhook"
    )


    await bot.set_webhook(
        webhook_url
    )


    print(
        "============================================"
    )

    print(
        "✅ BOT STARTED"
    )

    print(
        "Webhook:",
        webhook_url
    )

    print(
        "Предметов:",
        len(TESTS)
    )

    print(
        "Всего вопросов:",
        sum(
            len(q)
            for q in TESTS.values()
        )
    )

    print(
        "============================================"
    )


# ============================================================
# SHUTDOWN
# ============================================================

async def on_shutdown(app):

    try:

        await bot.delete_webhook()

    except Exception:
        pass


    try:

        await bot.session.close()

    except Exception:
        pass


# ============================================================
# WEB SERVER
# ============================================================

app = web.Application()


handler = SimpleRequestHandler(
    dispatcher=dp,
    bot=bot
)


handler.register(
    app,
    path="/webhook"
)


setup_application(
    app,
    dp,
    bot=bot
)


app.on_startup.append(
    on_startup
)


app.on_cleanup.append(
    on_shutdown
)


# ============================================================
# START
# ============================================================

if __name__ == "__main__":

    web.run_app(
        app,
        host="0.0.0.0",
        port=PORT
    )
